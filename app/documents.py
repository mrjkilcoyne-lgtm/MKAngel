"""
Document management for MKAngel.

Offline-first: documents save locally as Quill Delta JSON.
Export to DOCX/PDF when requested. Cloud sync when online.

Storage: ~/.mkangel/documents/
"""

from __future__ import annotations

import json
import os
import time
from dataclasses import dataclass, field
from pathlib import Path


from app.paths import mkangel_dir

DOCS_DIR = mkangel_dir() / "documents"


@dataclass
class Document:
    """A single document managed by the Angel."""
    doc_id: str
    title: str
    content: dict = field(default_factory=dict)  # Quill Delta JSON
    created_at: float = 0.0
    updated_at: float = 0.0
    version: int = 1
    tags: list[str] = field(default_factory=list)

    def __post_init__(self):
        if not self.created_at:
            self.created_at = time.time()
        if not self.updated_at:
            self.updated_at = time.time()


class DocumentManager:
    """Manages document lifecycle: create, save, load, list, export."""

    def __init__(self, docs_dir: Path | None = None):
        self.docs_dir = docs_dir or DOCS_DIR
        self.docs_dir.mkdir(parents=True, exist_ok=True)

    def create(self, title: str = "Untitled") -> Document:
        doc_id = f"doc_{int(time.time() * 1000)}"
        doc = Document(doc_id=doc_id, title=title)
        self.save(doc)
        return doc

    def save(self, doc: Document) -> None:
        doc.updated_at = time.time()
        doc.version += 1
        path = self.docs_dir / f"{doc.doc_id}.json"
        path.write_text(json.dumps({
            "doc_id": doc.doc_id,
            "title": doc.title,
            "content": doc.content,
            "created_at": doc.created_at,
            "updated_at": doc.updated_at,
            "version": doc.version,
            "tags": doc.tags,
        }, indent=2), encoding="utf-8")

    def load(self, doc_id: str) -> Document | None:
        path = self.docs_dir / f"{doc_id}.json"
        if not path.exists():
            return None
        data = json.loads(path.read_text(encoding="utf-8"))
        return Document(**data)

    def list_documents(self) -> list[Document]:
        docs = []
        for f in sorted(self.docs_dir.glob("doc_*.json"), reverse=True):
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                docs.append(Document(**data))
            except Exception:
                continue
        return docs

    def delete(self, doc_id: str) -> bool:
        path = self.docs_dir / f"{doc_id}.json"
        if path.exists():
            path.unlink()
            return True
        return False

    def export_text(self, doc: Document) -> str:
        """Export document content as plain text."""
        content = doc.content
        if isinstance(content, dict) and "ops" in content:
            return "".join(
                op.get("insert", "") for op in content["ops"]
                if isinstance(op.get("insert"), str)
            )
        return str(content)

    def export_docx(self, doc: Document, output_path: str) -> str | None:
        """Export to DOCX. Returns path or None if python-docx unavailable."""
        try:
            from docx import Document as DocxDocument
            d = DocxDocument()
            d.add_heading(doc.title, 0)
            d.add_paragraph(self.export_text(doc))
            d.save(output_path)
            return output_path
        except ImportError:
            return None
